from openai import OpenAI
import streamlit as st 
import os
from streamlit_chat import message as msg 
import docx
import io

# Uso de senha com Variável de Ambiente
client = OpenAI(
    api_key=os.environ['SENHA_OPEN_AI'],        # verificar nova senha OpenAI
) 
# ou pode usar >>> client = OpenAI(api_key = os.getenv("SENHA_OPEN_AI"))

# ou pip install python_dotenv >>>
# from dotenv import load_dotenv
# import os

# load_dotenv()
# api_key = os.environ['SENHA_OPEN_AI']

st.title("Chat com ChatGPT")
st.write("***")

if "hst_conversa" not in st.session_state:
    st.session_state.hst_conversa = []

pergunta = st.text_area("Digite a pergunta:")
btn_enviar_msg = st.button("Enviar mensagem")
if btn_enviar_msg:
    st.session_state.hst_conversa.append({"role": "user", "content": pergunta})
    retorno_openai = client.chat.completions.create(
        model = "gpt-4.1",
        messages = st.session_state.hst_conversa,
        max_tokens = 1000,
        n=1,
    )
    # Conforme código descrito nos docs da OpenAI >>> https://platform.openai.com/docs/api-reference/chat/create >>>
    # st.write(retorno_openai.choices[0].message.content)
    st.session_state.hst_conversa.append(
        {"role": "assistant",
         "content": retorno_openai.choices[0].message.content}
    )
    
if len(st.session_state.hst_conversa) > 0:
    for i in range(len(st.session_state.hst_conversa)):
        if i % 2 == 0:
            msg("Você: " + st.session_state.hst_conversa[i]['content'], is_user=True)
        else:
            msg("Resposta IA: " + st.session_state.hst_conversa[i]['content'])

if len(st.session_state.hst_conversa) > 0:
    btn_salvar = st.button("Salvar Conteúdo")
    if btn_salvar:
        trabalho = io.BytesIO()     # io usado como forma de salvar o conteúdo pelo streamlit
        documento = docx.Document()
        documento.add_heading("Conteúdo Gerado", level=1)
        # percorre todo o documento > perguntas e respostas
        for i in range(len(st.session_state.hst_conversa)):
            if i % 2 == 0:
                documento.add_heading("Pergunta", level=2)
                documento.add_paragraph(st.session_state.hst_conversa[i]['content'])
            else:
                documento.add_heading("Resposta", level=2)
                documento.add_paragraph(st.session_state.hst_conversa[i]['content'])

        documento.save(trabalho)
        st.download_button(label="Clique aqui para salvar o conteúdo",
                           data=trabalho,
                           file_name="",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        